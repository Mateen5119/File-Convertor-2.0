"""DOCX Editor widget – load, edit, and save Word documents.

Uses python-docx for reading/writing .docx files.  Falls back gracefully
when the library is not installed.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional

from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QToolBar,
    QTextEdit,
    QLabel,
    QFileDialog,
    QMessageBox,
    QSizePolicy,
)

# ---------------------------------------------------------------------------
# Graceful import
# ---------------------------------------------------------------------------
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Constants – Stitch design tokens
# ---------------------------------------------------------------------------
_SURFACE = "#131315"
_ON_SURFACE = "#e4e2e4"
_PRIMARY = "#adc6ff"
_TERTIARY = "#47e266"
_ERROR = "#ffb4ab"
_OUTLINE = "#8b90a0"


# ---------------------------------------------------------------------------
# Worker: load DOCX in background
# ---------------------------------------------------------------------------
class _LoadWorker(QThread):
    """Loads a .docx file and extracts paragraph text."""

    finished = pyqtSignal(list)  # list[str] paragraphs
    error = pyqtSignal(str)

    def __init__(self, path: str) -> None:
        super().__init__()
        self._path = path

    def run(self) -> None:
        try:
            doc = DocxDocument(self._path)
            paragraphs = [p.text for p in doc.paragraphs]
            self.finished.emit(paragraphs)
        except Exception as e:
            self.error.emit(str(e))


# ---------------------------------------------------------------------------
# Worker: save DOCX in background
# ---------------------------------------------------------------------------
class _SaveWorker(QThread):
    """Writes edited paragraphs back to a .docx file."""

    finished = pyqtSignal(str)  # saved path
    error = pyqtSignal(str)

    def __init__(self, paragraphs: list[str], save_path: str, template_path: Optional[str] = None) -> None:
        super().__init__()
        self._paragraphs = paragraphs
        self._save_path = save_path
        self._template_path = template_path

    def run(self) -> None:
        try:
            if self._template_path and os.path.exists(self._template_path):
                # Re-open original to preserve styles, then overwrite text
                doc = DocxDocument(self._template_path)
                existing = doc.paragraphs
                for i, text in enumerate(self._paragraphs):
                    if i < len(existing):
                        existing[i].text = text
                    else:
                        doc.add_paragraph(text)
                # Remove extra paragraphs from original if user deleted some
                # python-docx doesn't support deletion easily, so we clear text
                for i in range(len(self._paragraphs), len(existing)):
                    existing[i].text = ""
            else:
                doc = DocxDocument()
                for text in self._paragraphs:
                    doc.add_paragraph(text)
            doc.save(self._save_path)
            self.finished.emit(self._save_path)
        except Exception as e:
            self.error.emit(str(e))


# ---------------------------------------------------------------------------
# Main widget
# ---------------------------------------------------------------------------
class DocxEditorWidget(QWidget):
    """A paragraph-level DOCX editor with Open / Save toolbar."""

    def __init__(self, parent: Optional[QWidget] = None) -> None:
        super().__init__(parent)
        self.setObjectName("glassCard")
        self._current_path: Optional[str] = None
        self._workers: list[QThread] = []
        self._build_ui()

    # ---- UI Construction ---------------------------------------------------

    def _build_ui(self) -> None:
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)

        # Dependency-missing banner
        if not HAS_DOCX:
            banner = QLabel(
                "⚠️  python-docx is not installed. Run  pip install python-docx  to enable DOCX editing."
            )
            banner.setObjectName("labelError")
            banner.setAlignment(Qt.AlignmentFlag.AlignCenter)
            banner.setWordWrap(True)
            root.addWidget(banner)
            return

        # Toolbar
        self._toolbar = self._create_toolbar()
        root.addWidget(self._toolbar)

        # Editor area
        editor_container = QWidget()
        editor_container.setObjectName("glassCard")
        editor_layout = QVBoxLayout(editor_container)
        editor_layout.setContentsMargins(12, 12, 12, 12)

        self._text_edit = QTextEdit()
        self._text_edit.setObjectName("glassCard")
        self._text_edit.setPlaceholderText("Open a .docx file to begin editing…")
        self._text_edit.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding
        )
        self._text_edit.setStyleSheet(
            f"QTextEdit {{ color: {_ON_SURFACE}; background: {_SURFACE}; "
            f"border: 1px solid {_OUTLINE}; border-radius: 8px; padding: 8px; "
            f"font-size: 14px; }}"
        )
        editor_layout.addWidget(self._text_edit)

        # Status bar
        self._status = QLabel("")
        self._status.setObjectName("labelOutline")
        self._status.setAlignment(Qt.AlignmentFlag.AlignRight)
        editor_layout.addWidget(self._status)

        root.addWidget(editor_container)

    def _create_toolbar(self) -> QToolBar:
        tb = QToolBar("DOCX Editor")
        tb.setMovable(False)
        tb.setObjectName("glassCardRow")

        self._act_open = tb.addAction("📂 Open")
        self._act_open.triggered.connect(self._open_file)

        self._act_save = tb.addAction("💾 Save")
        self._act_save.setEnabled(False)
        self._act_save.triggered.connect(self._save_file)

        tb.addSeparator()

        self._word_count_label = QLabel("  Words: 0")
        self._word_count_label.setObjectName("labelOutline")
        tb.addWidget(self._word_count_label)

        return tb

    # ---- Actions -----------------------------------------------------------

    def _open_file(self) -> None:
        path, _ = QFileDialog.getOpenFileName(
            self, "Open DOCX", "", "Word Documents (*.docx)"
        )
        if not path:
            return
        self._load_docx(path)

    def _load_docx(self, path: str) -> None:
        self._current_path = path
        self._status.setText("Loading…")
        self._text_edit.setEnabled(False)

        worker = _LoadWorker(path)
        worker.finished.connect(self._on_load_finished)
        worker.error.connect(self._on_load_error)
        self._workers.append(worker)
        worker.start()

    def _on_load_finished(self, paragraphs: list[str]) -> None:
        self._text_edit.setEnabled(True)
        self._text_edit.setPlainText("\n".join(paragraphs))
        self._act_save.setEnabled(True)
        self._update_word_count()
        self._text_edit.textChanged.connect(self._update_word_count)
        name = Path(self._current_path).name if self._current_path else ""
        self._status.setText(f"Loaded  ✓  {name}")

    def _on_load_error(self, message: str) -> None:
        self._text_edit.setEnabled(True)
        QMessageBox.critical(self, "Load Error", f"Failed to open file:\n{message}")
        self._status.setText("Load failed")

    def _update_word_count(self) -> None:
        text = self._text_edit.toPlainText()
        words = len(text.split()) if text.strip() else 0
        self._word_count_label.setText(f"  Words: {words}")

    def _save_file(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "Save DOCX", self._current_path or "", "Word Documents (*.docx)"
        )
        if not path:
            return

        paragraphs = self._text_edit.toPlainText().split("\n")
        self._status.setText("Saving…")

        worker = _SaveWorker(paragraphs, path, self._current_path)
        worker.finished.connect(self._on_save_finished)
        worker.error.connect(self._on_save_error)
        self._workers.append(worker)
        worker.start()

    def _on_save_finished(self, path: str) -> None:
        self._current_path = path
        self._status.setText(f"Saved  ✓  {Path(path).name}")

    def _on_save_error(self, message: str) -> None:
        QMessageBox.critical(self, "Save Error", f"Failed to save file:\n{message}")
        self._status.setText("Save failed")

    # ---- Public API --------------------------------------------------------

    def load(self, path: str) -> None:
        """Programmatically open a DOCX file."""
        if HAS_DOCX:
            self._load_docx(path)
